"""
Mermaid图表转换器模块
"""
import os
import json
import time
import tempfile
import subprocess
from pathlib import Path
from docx.shared import Inches
from .base import ElementConverter


class MermaidConverter(ElementConverter):
    """Mermaid图表转换器，处理Markdown中的mermaid代码块"""

    def __init__(self, base_converter=None):
        """初始化Mermaid转换器
        
        Args:
            base_converter: 基础转换器实例
        """
        super().__init__(base_converter)
        self.debug = False
        self.max_retries = 3  # 最大重试次数
        self.retry_delay = 1  # 重试间隔（秒）
        if base_converter:
            self.debug = base_converter.debug

    def _generate_mermaid_image(self, mmd_file, config_file, img_file):
        """生成Mermaid图表图片
        
        Args:
            mmd_file: Mermaid代码文件路径
            config_file: 配置文件路径
            img_file: 输出图片文件路径
            
        Returns:
            bool: 是否成功生成图片
        """
        cmd = [
            'mmdc',
            '-i', mmd_file,
            '-o', img_file,
            '-c', config_file,
            '-b', 'transparent'
        ]
        
        if self.debug:
            print(f"执行命令: {' '.join(cmd)}")
            cmd.append('--verbose')

        for attempt in range(self.max_retries):
            try:
                if self.debug:
                    print(f"尝试生成图片，第 {attempt + 1} 次")
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)  # 添加超时限制
                
                if result.returncode == 0 and os.path.exists(img_file):
                    if self.debug:
                        print(f"图片生成成功: {img_file}")
                        print(f"图片大小: {os.path.getsize(img_file)} 字节")
                    return True
                
                if self.debug:
                    print(f"命令执行失败 (尝试 {attempt + 1}/{self.max_retries})")
                    print(f"错误输出: {result.stderr}")
                    print(f"标准输出: {result.stdout}")
                
                if attempt < self.max_retries - 1:
                    time.sleep(self.retry_delay)
                    
            except subprocess.TimeoutExpired:
                if self.debug:
                    print(f"命令执行超时 (尝试 {attempt + 1}/{self.max_retries})")
                if attempt < self.max_retries - 1:
                    time.sleep(self.retry_delay)
            except Exception as e:
                if self.debug:
                    print(f"生成图片时发生异常: {str(e)}")
                if attempt < self.max_retries - 1:
                    time.sleep(self.retry_delay)
        
        return False

    def convert(self, token):
        """转换mermaid代码块为图片
        
        Args:
            token: mermaid代码块token
            
        Returns:
            docx.paragraph: 包含图片的段落
        """
        if not self.document:
            raise ValueError("Document not set for MermaidConverter")

        if self.debug:
            print(f"处理Mermaid图表: {token}")

        # 获取mermaid代码
        code = token.content if hasattr(token, 'content') else ''
        if not code:
            if self.debug:
                print("Mermaid代码为空")
            return None

        temp_dir = None
        try:
            # 创建临时目录
            temp_dir = tempfile.mkdtemp()
            if self.debug:
                print(f"创建临时目录: {temp_dir}")
            
            # 创建临时文件保存mermaid代码
            mmd_file = os.path.join(temp_dir, 'diagram.mmd')
            with open(mmd_file, 'w', encoding='utf-8') as f:
                f.write(code)
            if self.debug:
                print(f"已保存Mermaid代码到文件: {mmd_file}")

            # 创建配置文件
            config = {
                "theme": "default",
                "themeVariables": {
                    "fontSize": "16px",
                    "fontFamily": "arial",
                },
                "flowchart": {
                    "htmlLabels": True,
                    "curve": "linear"
                },
                "sequence": {
                    "showSequenceNumbers": False,
                    "actorMargin": 50,
                    "messageMargin": 40
                },
                "gantt": {
                    "leftPadding": 75,
                    "rightPadding": 20
                }
            }
            
            config_file = os.path.join(temp_dir, 'config.json')
            with open(config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f)

            # 生成图片文件路径
            img_file = os.path.join(temp_dir, 'diagram.png')
            if self.debug:
                print(f"目标图片文件路径: {img_file}")

            # 生成图片
            if not self._generate_mermaid_image(mmd_file, config_file, img_file):
                if self.debug:
                    print("图片生成失败，所有重试都失败了")
                return None

            # 创建新段落
            paragraph = self.document.add_paragraph()
            paragraph.alignment = 1  # 居中对齐

            # 添加图片到段落
            run = paragraph.add_run()
            run.add_picture(img_file, width=Inches(6))  # 设置合适的宽度

            return paragraph

        except Exception as e:
            if self.debug:
                print(f"Mermaid转换异常: {str(e)}")
            return None
            
        finally:
            # 确保在所有情况下都清理临时文件和目录
            if temp_dir and os.path.exists(temp_dir):
                try:
                    for file_name in os.listdir(temp_dir):
                        file_path = os.path.join(temp_dir, file_name)
                        if os.path.exists(file_path):
                            if self.debug:
                                print(f"删除临时文件: {file_path}")
                            os.unlink(file_path)
                    os.rmdir(temp_dir)
                    if self.debug:
                        print(f"删除临时目录: {temp_dir}")
                except Exception as e:
                    if self.debug:
                        print(f"清理临时文件时出错: {str(e)}")